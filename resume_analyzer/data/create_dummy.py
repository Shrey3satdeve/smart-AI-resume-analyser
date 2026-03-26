from docx import Document  # type: ignore
import os

os.makedirs("d:/resume_analyzer/resume_analyzer/data", exist_ok=True)

# Resume 1
doc1 = Document()
doc1.add_heading('John ML Doe', 0)
doc1.add_paragraph('john.doe@email.com | 111-222-3333')
doc1.add_paragraph('Skills: Python, Machine Learning, TensorFlow, AWS, SQL, Deep Learning, Data Science')
doc1.add_paragraph('Experience: Data Scientist at Tech Corp building cool AI models.')
doc1.add_paragraph('Education: BS in Computer Science from University of AI.')
doc1.save("d:/resume_analyzer/resume_analyzer/data/test_resume_1.docx")

# Resume 2
doc2 = Document()
doc2.add_heading('Jane Web Smith', 0)
doc2.add_paragraph('jane.smith@email.com | 987-654-3210')
doc2.add_paragraph('Skills: React, Node.js, JavaScript, HTML, CSS, MongoDB, Git, TypeScript, Frontend')
doc2.add_paragraph('Experience: Frontend dev at WebWorks for 3 years.')
doc2.save("d:/resume_analyzer/resume_analyzer/data/test_resume_2.docx")

print("Created 2 dummy docx resumes.")
